import io
from typing import Dict, Any
import logging

logger = logging.getLogger(__name__)

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

class DOCXGenerator:
    @staticmethod
    def generate(title: str, analytics: Dict[str, Any], forecast: Dict[str, Any], recommendations: Dict[str, Any]) -> bytes:
        if not HAS_DOCX:
            # Fallback simple text formatting for docx binary stream mock
            logger.warning("python-docx is not installed. Generating simulated word docx binary stream.")
            buffer = io.BytesIO()
            buffer.write(f"=== {title} ===\n".encode("utf-8"))
            buffer.write(f"KPIs: {str(analytics.get('kpis'))}\n".encode("utf-8"))
            buffer.write(f"Forecast: {str(forecast.get('trend'))}\n".encode("utf-8"))
            return buffer.getvalue()
            
        doc = Document()
        doc.add_heading(title, 0)
        
        # Summary
        doc.add_heading("1. Executive Summary", level=1)
        doc.add_paragraph("Compiled business intelligence reports detailing corporate KPI performance and predictions.")
        
        # KPIs
        doc.add_heading("2. KPI Dashboard", level=1)
        for k, v in analytics.get("kpis", {}).items():
            doc.add_paragraph(f"{k}: {v}")
            
        # Forecast
        doc.add_heading("3. Forecasting Results", level=1)
        doc.add_paragraph(f"Trend Direction: {forecast.get('trend', 'Flat')}")
        doc.add_paragraph(f"Model Used: {forecast.get('model_used', 'ARIMA')}")
        
        # Recs
        doc.add_heading("4. Prescriptive Recommendations", level=1)
        for r in recommendations.get("recommendations", []):
            doc.add_paragraph(f"[{r.get('priority')}] {r.get('title')}: {', '.join(r.get('suggested_actions', []))}")
            
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
